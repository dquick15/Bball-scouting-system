"""
modules/report_gen/report_generator.py
========================================
OpenAI-powered scouting report generator + DOCX/PDF export helpers.
Import path updated to use the local prompts module.
"""
from __future__ import annotations

import io
import os
from textwrap import wrap

from docx import Document
from openai import OpenAI
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from modules.report_gen.prompts import build_system_prompt, build_user_prompt


class ScoutingReportGenerator:
    def __init__(self, api_key: str | None = None, model: str | None = None) -> None:
        resolved_key = api_key or os.getenv("OPENAI_API_KEY")
        if not resolved_key:
            raise ValueError(
                "OpenAI API key not found. Set OPENAI_API_KEY or add it to Streamlit secrets."
            )
        self.client = OpenAI(api_key=resolved_key)
        self.model = model or os.getenv("OPENAI_MODEL", "gpt-4.1-mini")

    def generate_report(self, player_record: dict[str, object], mode: str) -> str:
        response = self.client.responses.create(
            model=self.model,
            input=[
                {
                    "role": "system",
                    "content": [{"type": "input_text", "text": build_system_prompt()}],
                },
                {
                    "role": "user",
                    "content": [{"type": "input_text", "text": build_user_prompt(player_record, mode)}],
                },
            ],
            max_output_tokens=500,
        )
        report_text = response.output_text.strip()
        if not report_text:
            raise ValueError("The OpenAI response did not contain report text.")
        return report_text


def create_docx_bytes(player_name: str, mode: str, report_text: str) -> bytes:
    document = Document()
    document.add_heading(f"Scouting Report: {player_name}", level=1)
    document.add_paragraph(f"Mode: {mode}")
    document.add_paragraph(report_text)
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def create_pdf_bytes(player_name: str, mode: str, report_text: str) -> bytes:
    buffer = io.BytesIO()
    document = SimpleDocTemplate(
        buffer, pagesize=letter,
        topMargin=54, bottomMargin=54, leftMargin=54, rightMargin=54,
    )
    styles = getSampleStyleSheet()
    story = [
        Paragraph(f"Scouting Report: {player_name}", styles["Title"]),
        Spacer(1, 12),
        Paragraph(f"Mode: {mode}", styles["Heading3"]),
        Spacer(1, 12),
    ]
    for paragraph in report_text.split("\n"):
        cleaned = paragraph.strip()
        if cleaned:
            story.append(Paragraph(cleaned.replace("\n", " "), styles["BodyText"]))
            story.append(Spacer(1, 10))
    if len(story) == 4:
        story.append(Paragraph("<br/>".join(wrap(report_text, width=95)), styles["BodyText"]))
    document.build(story)
    buffer.seek(0)
    return buffer.getvalue()
